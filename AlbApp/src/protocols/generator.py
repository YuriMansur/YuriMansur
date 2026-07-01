"""Генератор протокола испытания в .docx (python-docx).

Строит готовый документ из данных, собранных мастером испытания (секция 3):
ГОСТ, методика, образец, зарегистрированные значения по шагам и заключение.
Сохраняет в protocols/out/. Форматирование задаётся кодом (таблица с рамками),
поэтому шаблон не требуется — позже можно перейти на docxtpl с .docx-шаблоном.
"""
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# все документы (протоколы) складываем в папку проекта documents/
DOCUMENTS_DIR = Path(__file__).parent.parent.parent / "documents"
_OUT_DIR = DOCUMENTS_DIR

# служебные ключи form-значений мастера — в протокол не выводим
_SKIP_PREFIX = "__"


def _collect_rows(fields: dict) -> list[tuple[str, str, str]]:
    """Развернуть form-значения мастера в строки (шаг, параметр, значение)."""
    rows = []
    for step_id, vals in (fields or {}).items():
        if not isinstance(vals, dict):
            continue
        for key, val in vals.items():
            if key.startswith(_SKIP_PREFIX):
                continue
            if val is None or val == "" or isinstance(val, bool):
                continue
            rows.append((str(step_id), str(key), str(val)))
    return rows


def _safe_name(name: str) -> str:
    """Убрать из имени папки символы, недопустимые в путях Windows."""
    for ch in '<>:"/\\|?*':
        name = name.replace(ch, "_")
    return name.strip().rstrip(".") or "Испытание"


def _write_header(doc: "Document", data: dict) -> None:
    doc.add_paragraph(f"ГОСТ: {data.get('gost', '')}")
    doc.add_paragraph(
        f"Методика: {data.get('method', '')} — {data.get('title', '')}")
    if data.get("stand"):
        doc.add_paragraph(f"Стенд: {data['stand']}")


def _write_sample_info(doc: "Document", data: dict) -> None:
    sample_rows = [
        ("Название образца", data.get("sample_name")),
        ("Дата получения образца", data.get("date_received")),
        ("Дата проведения", data.get("date_conducted")),
        ("Образец используется", data.get("used")),
        ("Используется взамен разрушенного", data.get("used_replacement")),
        ("Применяемые концевые крепления", data.get("clamps")),
    ]
    if any(v for _, v in sample_rows):
        doc.add_heading("Информация об исследуемом образце", level=1)
        for label, val in sample_rows:
            doc.add_paragraph(f"{label}: {val or '—'}")


def _write_values_table(doc: "Document", rows: list, heading: str) -> None:
    doc.add_heading(heading, level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Шаг", "Параметр", "Значение"
    for step, key, val in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = step, key, val


def _build_protocol(data: dict) -> "Document":
    doc = Document()
    title = doc.add_heading("ПРОТОКОЛ ИСПЫТАНИЯ", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _write_header(doc, data)
    _write_sample_info(doc, data)
    doc.add_paragraph(f"Дата формирования: {datetime.now():%d.%m.%Y %H:%M}")

    rows = _collect_rows(data.get("fields"))
    if rows:
        _write_values_table(doc, rows, "Зарегистрированные значения")

    passed = data.get("passed")
    if passed is not None:
        doc.add_heading("Заключение", level=1)
        doc.add_paragraph(
            "Образец ВЫДЕРЖАЛ испытание." if passed
            else "Образец НЕ выдержал требования испытания.")
    return doc


def _build_journal(data: dict) -> "Document":
    doc = Document()
    title = doc.add_heading("ЖУРНАЛ ИСПЫТАНИЯ", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _write_header(doc, data)
    doc.add_paragraph(f"Дата формирования: {datetime.now():%d.%m.%Y %H:%M}")

    rows = _collect_rows(data.get("fields"))
    if rows:
        _write_values_table(doc, rows, "Ход испытания")
    else:
        doc.add_paragraph("Зарегистрированных значений нет.")
    return doc


def generate_test_docs(data: dict, out_dir=None) -> Path:
    """Создать папку испытания и сохранить в неё Протокол.docx и Журнал.docx.

    Имя папки — номер испытания (data['test_no']), иначе методика + дата/время.
    Возвращает путь к созданной папке.
    """
    base = Path(out_dir) if out_dir else _OUT_DIR
    num = str(data.get("test_no") or "").strip()
    if not num:
        method = str(data.get("method", "")).replace("/", "_").replace(".", "_")
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        num = f"{method}_{stamp}" if method else stamp

    folder = base / _safe_name(num)
    folder.mkdir(parents=True, exist_ok=True)
    _build_protocol(data).save(folder / "Протокол.docx")
    _build_journal(data).save(folder / "Журнал.docx")
    return folder
